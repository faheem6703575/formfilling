from utils.ai_helper import AIHelper
from docx import Document
import tempfile
import os
from datetime import datetime
import re
from io import BytesIO

class ComprehensiveDocumentFillingAgent:
    def __init__(self):
        self.ai_helper = AIHelper()
    
    def fill_declaration_form(self, uploaded_file, comprehensive_data):
        """Fill Document 1: Declaration Form"""
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        try:
            doc = Document(tmp_path)
            
            # Extract data from comprehensive dataset
            company_name = comprehensive_data.get('COMPANY_NAME', 'Baltic Innovation Technologies UAB')
            company_code = comprehensive_data.get('COMPANY_CODE', 'LT305123456')
            completion_date = comprehensive_data.get('COMPLETION_DATE', datetime.now().strftime("%d.%m.%Y"))
            manager_name = comprehensive_data.get('MANAGER_NAME', 'Dr. Vytautas Kazlauskas')
            manager_position = comprehensive_data.get('MANAGER_POSITION', 'CEO')
            
            # Fill the date at the top
            date_filled = False
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if "date of completion" in text.lower():
                    new_text = re.sub(r'_{3,}', completion_date, text)
                    if new_text != text:
                        paragraph.text = new_text
                        date_filled = True
                        break
            
            # If not found, look for any paragraph with multiple underscores near the top
            if not date_filled:
                for i, paragraph in enumerate(doc.paragraphs[:10]):
                    text = paragraph.text
                    if re.search(r'_{10,}', text) and len(text.strip()) < 100:
                        new_text = re.sub(r'_{10,}', completion_date, text)
                        paragraph.text = new_text
                        date_filled = True
                        break
            
            # Fill the table with company information
            for table in doc.tables:
                rows_list = list(table.rows)
                
                for row_idx, row in enumerate(rows_list):
                    row_text = ' '.join([cell.text for cell in row.cells]).strip()
                    
                    # Company name field
                    if "Name of the declaring undertaking" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            for cell in next_row.cells:
                                if cell.text.strip() == "" or len(cell.text.strip()) < 5:
                                    cell.text = "      " + company_name
                                    break
                    
                    # Company code field
                    elif "Code of the declaring company" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            for cell in next_row.cells:
                                if cell.text.strip() == "" or len(cell.text.strip()) < 5:
                                    cell.text = "      " + company_code
                                    break
            
            # Fill signature area
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if "position of the manager" in text.lower() or ("signature" in text.lower() and "_" in text):
                    signature_line = f"{manager_position}                                    {manager_name}"
                    new_text = re.sub(r'_{20,}.*?_{10,}', signature_line, text)
                    if new_text != text:
                        paragraph.text = new_text
            
            # Save to BytesIO
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return {
                "success": True,
                "filled_document": doc_bytes,
                "filename": f"filled_declaration_{company_name.replace(' ', '_')}.docx",
                "preview": f"Declaration filled for {company_name} (Code: {company_code})"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error filling declaration document: {str(e)}"
            }
        finally:
            try:
                os.unlink(tmp_path)
            except:
                pass
    
    def fill_research_areas_in_document(self, doc, comprehensive_data):
        """Fill research areas in the MTEP document"""
        research_area = comprehensive_data.get('RESEARCH_AREA', 'Computer Sciences')
        
        # Look for research area table
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                row_text = ' '.join([cell.text for cell in row.cells]).strip()
                
                # Find the research areas table header
                if "Research area(s) of the project" in row_text and "Indicate up to two areas" in row_text:
                    if row_idx + 1 < len(table.rows):
                        content_row = table.rows[row_idx + 1]
                        if len(content_row.cells) >= 1:
                            content_row.cells[0].text = f"1. {research_area}(main)\n2. Artificial Intelligence"
                            break
                
                # Look for cells that contain research area patterns
                for cell in row.cells:
                    if "Research area(s) of the project" in cell.text:
                        cell_text = cell.text.strip()
                        if "1." not in cell_text:
                            cell_idx = list(row.cells).index(cell)
                            if cell_idx + 1 < len(row.cells):
                                row.cells[cell_idx + 1].text = f"1. {research_area}(main)\n2. Artificial Intelligence"
                                return
                            if row_idx + 1 < len(table.rows):
                                next_row = table.rows[row_idx + 1]
                                if len(next_row.cells) >= 1:
                                    next_row.cells[0].text = f"1. {research_area}(main)\n2. Artificial Intelligence"
                                    return

    def add_text_to_paragraph(self, paragraph, text_to_add):
        """Helper method to add text to an existing paragraph"""
        if paragraph.text.strip():
            paragraph.text = paragraph.text + "\n\n" + text_to_add
        else:
            paragraph.text = text_to_add

    def find_and_fill_summary_section(self, doc, summary):
        """Find and fill the summary section in the document"""
        summary_filled = False
        
        # Method 1: Look for specific summary patterns in paragraphs
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.lower().strip()
            
            # Look for summary-related text patterns
            if any(pattern in para_text for pattern in [
                "what are the long-term objectives set",
                "long-term objectives",
                "project summary",
                "summary of the project",
                "describe the project"
            ]):
                # Try to find the next empty or nearly empty paragraph
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    if len(next_para.text.strip()) < 20:  # Empty or very short paragraph
                        next_para.text = summary
                        summary_filled = True
                        break
                
                if summary_filled:
                    break
                
                # If no empty paragraph found, add to current paragraph
                self.add_text_to_paragraph(para, f"\n\n{summary}")
                summary_filled = True
                break
        
        # Method 2: If not found in paragraphs, look in table cells
        if not summary_filled:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.lower().strip()
                        if any(pattern in cell_text for pattern in [
                            "objectives",
                            "summary",
                            "description",
                            "project goals"
                        ]):
                            if len(cell.text.strip()) < 50:  # Relatively empty cell
                                cell.text = summary
                                summary_filled = True
                                break
                    if summary_filled:
                        break
                if summary_filled:
                    break
        
        # Method 3: If still not found, add to first suitable paragraph
        if not summary_filled:
            for i, para in enumerate(doc.paragraphs):
                if len(para.text.strip()) < 10 and i > 5:  # Skip header paragraphs
                    para.text = f"Project Summary:\n{summary}"
                    summary_filled = True
                    break
        
        return summary_filled
    
    def find_and_fill_innovativeness_section(self, doc, innovativeness):
        """Find and fill the innovativeness section in the document (Section 3.1.1)"""
        
        # WRITE YOUR INNOVATIVENESS FILLING LOGIC BELOW:
        innovativeness_filled = False
        
        # Method 1: Find the exact text "new products will bring to consumers, what problems they will solve, etc." and fill below it
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.lower().strip()
            
            # Look for the specific ending text pattern
            if "new products will bring to consumers, what problems they will solve" in para_text:
                # Found the target paragraph, now fill the next empty paragraph or create content below
                
                # Try to find the next empty or nearly empty paragraph
                content_filled = False
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    if len(next_para.text.strip()) < 30:  # Empty or very short paragraph
                        next_para.text = innovativeness
                        innovativeness_filled = True
                        content_filled = True
                        break
                
                # If no empty paragraph found, add content directly to the current paragraph
                if not content_filled:
                    if para.text.strip().endswith('.') or para.text.strip().endswith('etc.'):
                        para.text = para.text + "\n\n" + innovativeness
                    else:
                        para.text = para.text + " " + innovativeness
                    innovativeness_filled = True
                break
            
        return innovativeness_filled

    def find_and_fill_ipr_section(self, doc, ipr):
        """Find and fill the IPR section in the document (Section 3.1.4)"""

        ipr_filled = False

        # Method 1: Find the exact text "Intellectual Property Rights" and fill below it
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.lower().strip()

            if "patenting (if patenting activities are included in the scope of the project)" in para_text:
                # Found the target paragraph, now fill the next empty paragraph or create content below

                # Try to find the next empty or nearly empty paragraph
                content_filled = False
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    if len(next_para.text.strip()) < 30:  # Empty or very short paragraph
                        next_para.text = ipr
                        ipr_filled = True
                        content_filled = True
                        break

                # If no empty paragraph found, add content directly to the current paragraph
                if not content_filled:
                    if para.text.strip().endswith('.') or para.text.strip().endswith('etc.'):
                        para.text = para.text + "\n\n" + ipr
                    else:
                        para.text = para.text + " " + ipr
                    ipr_filled = True
                break

        return ipr_filled

    def find_and_fill_keywords_section(self, doc, keywords):
        """Find and fill the keywords section in the document (Section 3.1.5)"""
        
        # WRITE YOUR KEYWORDS FILLING LOGIC BELOW:
        keywords_filled = False
        
        # Method 1: Look for specific keywords patterns in paragraphs
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.lower().strip()
            
            # Look for keywords-related text patterns
            if "keywords" in para_text or "key words" in para_text:
                # Try to find the next empty or nearly empty paragraph
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    if len(next_para.text.strip()) < 20:
                        next_para.text = keywords
                        keywords_filled = True
                        break
                
                if not keywords_filled:
                    if para.text.strip().endswith('.'):
                        para.text = para.text + "\n\n" + keywords
                    else:
                        para.text = para.text + " " + keywords
                    keywords_filled = True
                break
        return keywords_filled
    
    def find_and_fill_commercialization_section(self, doc, commercialization):
        """Find and fill the commercialization section in the document (Section 4.7)"""
        
        commercialization_filled = False
        
        # Method 1: Look for specific commercialization patterns in paragraphs
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.lower().strip()
            
            # Look for commercialization-related text patterns
            if "to be completed if the product market development activities are included in the scope of the project" in para_text or "marketing approach" in para_text:
                # Try to find the next empty or nearly empty paragraph
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    if len(next_para.text.strip()) < 20:
                        next_para.text = commercialization
                        commercialization_filled = True
                        break
                
                if not commercialization_filled:
                    if para.text.strip().endswith('.'):
                        para.text = para.text + "\n\n" + commercialization
                    else:
                        para.text = para.text + " " + commercialization
                    commercialization_filled = True
                break
        
        return commercialization_filled
    
    def find_and_fill_literature_review_section(self, doc, literature_review):
        """Find and fill the literature review section in the document (Section 4.4)"""

        literature_review_filled = False

        # Method 1: Look for specific literature review patterns in paragraphs
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.lower().strip()

            # Look for literature review-related text patterns
            if "the research, up to 3 pages" in para_text:
                # Try to find the next empty or nearly empty paragraph
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    if len(next_para.text.strip()) < 20:
                        next_para.text = literature_review
                        literature_review_filled = True
                        break

                if not literature_review_filled:
                    if para.text.strip().endswith('.'):
                        para.text = para.text + "\n\n" + literature_review
                    else:
                        para.text = para.text + " " + literature_review
                    literature_review_filled = True
                break

        return literature_review_filled

    def find_and_fill_collaboration_section(self, doc, collaboration):
        """Find and fill the collaboration section in the document (Section 4.9)"""

        collaboration_filled = False

        # Method 1: Look for specific collaboration patterns in paragraphs
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.lower().strip()

            # Look for collaboration-related text patterns
            if "implemented with partners" in para_text or "partnership" in para_text:
                # Try to find the next empty or nearly empty paragraph
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    if len(next_para.text.strip()) < 20:
                        next_para.text = collaboration
                        collaboration_filled = True
                        break

                if not collaboration_filled:
                    if para.text.strip().endswith('.'):
                        para.text = para.text + "\n\n" + collaboration
                    else:
                        para.text = para.text + " " + collaboration
                    collaboration_filled = True
                break

        return collaboration_filled


    def find_and_fill_literature_sources_section(self, doc, literature_sources):
        """Find and fill the literature sources section in the document (Section 8)"""

        literature_sources_filled = False

        # Method 1: Look for specific literature sources patterns in paragraphs
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.lower().strip()

            # Look for literature sources-related text patterns
            if "Literature sources used" in para_text:
                # Try to find the next empty or nearly empty paragraph
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    if len(next_para.text.strip()) < 20:
                        next_para.text = literature_sources
                        literature_sources_filled = True
                        break

                if not literature_sources_filled:
                    if para.text.strip().endswith('.'):
                        para.text = para.text + "\n\n" + literature_sources
                    else:
                        para.text = para.text + " " + literature_sources
                    literature_sources_filled = True
                break

        return literature_sources_filled

    def fill_mtep_business_plan(self, uploaded_file, comprehensive_data):
        """Fill Document 2: MTEP Business Plan - FIXED"""
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        try:
            doc = Document(tmp_path)
            
            # Fill research areas first
            self.fill_research_areas_in_document(doc, comprehensive_data)
            
            # Extract data from comprehensive dataset and convert to strings
            company_name = str(comprehensive_data.get('COMPANY_NAME', 'Baltic Innovation Technologies UAB'))
            Shareholding_Sh = str(comprehensive_data.get('Sharehol', '50'))
            summary = str(comprehensive_data.get('SUMMARY_1', 'This section provides a summary of the applicant’s identity, main activities, and the project’s goals, including development, commercialization potential, and funding needs. It highlights the product’s uniqueness, target market, added value, required R&D and patenting costs, and long-term objectives.'))
            innovativeness = str(comprehensive_data.get('INNOVATIVENESS', 'Our product features advanced AI algorithms with real-time processing capabilities, offering 95% accuracy compared to 70% of existing solutions. The innovative machine learning architecture provides personalized recommendations, solving the problem of generic one-size-fits-all approaches. This brings significant benefits to consumers through improved efficiency, cost reduction of up to 40%, and enhanced user experience with intuitive interfaces.'))
            keywords = str(comprehensive_data.get('PROJECT_KEYWORDS', 'AI, machine learning, real-time processing, personalized recommendations, efficiency, cost reduction'))
            ipr = str(comprehensive_data.get('IPR', 'The product will be patented in the EU and US. The intellectual property will be owned by Baltic Innovation Technologies UAB. Patent justification includes unique AI algorithms and data processing methods that are not currently available in the market.'))
            commercialization = str(comprehensive_data.get('COMMERCIALIZATION', 'The product will target healthcare providers in the EU and US markets, leveraging online marketing and partnerships with medical institutions. Sales strategy includes direct sales, online platforms, and participation in industry conferences.'))
            literature_review = str(comprehensive_data.get('LITERATURE_REVIEW', 'Recent studies indicate a significant gap in AI-powered diagnostic systems, with most existing solutions lacking real-time processing capabilities. Our literature review covers over 100 sources, highlighting the need for more personalized and efficient healthcare solutions.'))
            collaboration = str(comprehensive_data.get('COLLABORATION', 'The project will be implemented in collaboration with TechCorp Europe, leveraging their expertise in AI and machine learning. This partnership will enhance the product’s development and market reach, providing access to advanced technologies and research facilities.'))
            literature_sources = str(comprehensive_data.get('LITERATURE_SOURCES', 'Smith et al. (2022), Johnson (2021), AI Innovations Journal (2023)'))
            product_name = str(comprehensive_data.get('PRODUCT_NAME', 'AI-powered diagnostic system'))
            current_tpl = str(comprehensive_data.get('CURRENT_TPL', 'TPL 3'))
            target_tpl = str(comprehensive_data.get('TARGET_TPL', 'TPL 6'))
            company_code_2 = str(comprehensive_data.get('COMPANY_CODE', 'LT305123456'))
            S_Ss_2 = str(comprehensive_data.get('S_S', '50'))
            S_Is_2 = str(comprehensive_data.get('S_I', 'MT23456'))
            S_Hs_2 = str(comprehensive_data.get('S_H', '50'))
            Name_of_the_legal_entity_2 = str(comprehensive_data.get('N_L_E', company_name))
            Identification_code_2 = str(comprehensive_data.get('I_C', 'LT123456789'))
            A_S_Ns_2 = str(comprehensive_data.get('A_S_Ns', 'Vytautas Kazlauskas'))
            share_Hs_2 = str(comprehensive_data.get('SHARE_HS', '100'))
            main_activity_2 = str(comprehensive_data.get('MAIN_ACTIVITY', 'Software development and innovation'))
            activity_percentage_2 = str(comprehensive_data.get('ACTIVITY_PERCENTAGE', '75'))
            cese_class_2 = str(comprehensive_data.get('CESE_CLASS', '62.01'))
            Pro_off_2 = str(comprehensive_data.get('PRODUCTS_OFFERED', 'AI-powered diagnostic system'))
            Per_sales_2 = str(comprehensive_data.get('PER_SALES', '100'))
            product_name_2 = str(comprehensive_data.get('PRODUCT_NAME', 'AI-powered diagnostic system'))
            novelty_level_2 = str(comprehensive_data.get('NOVELTY_LEVEL', 'market level'))
            jus_pro_2 = str(comprehensive_data.get('JUS_PRO', 'This product introduces new technological solutions not available in the current market'))
            rd_priority_2 = str(comprehensive_data.get('RD_PRIORITY', 'Information and communication technologies'))
            jus_r_d_i_2 = str(comprehensive_data.get('JUS_R_D_I', 'This project aligns with smart specialization priorities in ICT and digital innovation'))
            N_As_2 = str(comprehensive_data.get('N_As', 'R&D Laboratory'))
            F_Os_2 = str(comprehensive_data.get('F_Os', 'Owned'))
            S_Us_2 = str(comprehensive_data.get('S_Us', '200 m²'))
            W_R_Ds_2 = str(comprehensive_data.get('W_R_Ds', 'All R&D activities, testing, prototype development'))            
            competitor_2 = str(comprehensive_data.get('COMPETITOR_M', 'TechCorp Europe'))
            competitor_market_share_2 = str(comprehensive_data.get('COMPETITOR_MARKET_SHARE', '25'))
            current_tpl_2 = str(comprehensive_data.get('CURRENT_TPL', 'TPL 3'))
            target_tpl_2 = str(comprehensive_data.get('TARGET_TPL', 'TPL 6'))
            revenue_projection_2 = str(comprehensive_data.get('REVENUE_PROJECTION', '500000'))
            rd_budget_2 = str(comprehensive_data.get('RD_BUDGET', '200000'))
            revenue_ratio_2 = str(comprehensive_data.get('REVENUE_RATIO', '2.5'))
            project_impact_title_2 = str(comprehensive_data.get('PROJECT_IMPACT_TITLE', f'Development of {product_name_2}'))
            project_start_month_2 = str(comprehensive_data.get('PROJECT_START_MONTH', '1'))
            project_completion_month_2 = str(comprehensive_data.get('PROJECT_COMPLETION_MONTH', '24'))
            E_s_RES_2 = str(comprehensive_data.get('E_S_RES', 'NOT_FOUND'))
            E_s_R_D_2 = str(comprehensive_data.get('E_S_R&D', 'NOT_FOUND'))
            E_s_R_2 = str(comprehensive_data.get('E_S_R', 'NOT_FOUND'))
            A_s_RES_2 = str(comprehensive_data.get('A_S_RES', 'NOT_FOUND'))
            A_s_R_D_2 = str(comprehensive_data.get('A_S_R&D', 'NOT_FOUND'))
            A_s_R_2 = str(comprehensive_data.get('A_S_R', 'NOT_FOUND'))
            A_s_P_2 = str(comprehensive_data.get('A_S_P', 'NOT_FOUND'))

            N_e_2 = str(comprehensive_data.get('N_E', 'NOT_FOUND'))
            N_r_2 = str(comprehensive_data.get('N_R', 'NOT_FOUND'))
            N_t_2 = str(comprehensive_data.get('N_T', 'NOT_FOUND'))
            N_w_T_2 = str(comprehensive_data.get('N_W_T', 'NOT_FOUND'))
            N_p_T_2 = str(comprehensive_data.get('N_P_T', 'NOT_FOUND'))



            project_impact_description = str(comprehensive_data.get('PROJECT_IMPACT_DESCRIPTION', 
                f'The objective is to develop {product_name_2} through systematic R&D activities.'))
            tpl_justification = str(comprehensive_data.get('TPL_JUSTIFICATION', 
                f'Current {current_tpl_2} achieved through laboratory validation. Target {target_tpl_2} through systematic development.'))
            
            # Clean percentage values
            activity_percentage_2 = activity_percentage_2.replace('%', '').strip()
            competitor_market_share_2 = competitor_market_share_2.replace('%', '').strip()            
            # Extract Risk Assessment data
            risk_stages = [
                (
                    str(comprehensive_data.get('RISK_STAGE_1', 'Concept formulation and feasibility validation')),
                    str(comprehensive_data.get('RISK_DESCRIPTION_1', 'Market acceptance uncertainty')),
                    str(comprehensive_data.get('CRITICAL_POINT_1', 'User needs validation')),
                    str(comprehensive_data.get('MITIGATION_ACTION_1', 'Extensive market research'))
                ),
                (
                    str(comprehensive_data.get('RISK_STAGE_2', 'Layout development, testing, checking')),
                    str(comprehensive_data.get('RISK_DESCRIPTION_2', 'Technical complexity')),
                    str(comprehensive_data.get('CRITICAL_POINT_2', 'Algorithm optimization')),
                    str(comprehensive_data.get('MITIGATION_ACTION_2', 'Iterative testing approach'))
                ),
                (
                    str(comprehensive_data.get('RISK_STAGE_3', 'Prototype development and demonstration')),
                    str(comprehensive_data.get('RISK_DESCRIPTION_3', 'Integration challenges')),
                    str(comprehensive_data.get('CRITICAL_POINT_3', 'System compatibility')),
                    str(comprehensive_data.get('MITIGATION_ACTION_3', 'Modular development'))
                ),
                (
                    str(comprehensive_data.get('RISK_STAGE_4', 'Production and evaluation of a pilot batch')),
                    str(comprehensive_data.get('RISK_DESCRIPTION_4', 'Scale-up difficulties')),
                    str(comprehensive_data.get('CRITICAL_POINT_4', 'Performance at scale')),
                    str(comprehensive_data.get('', 'Pilot testing program'))
                )
            ]
            
            # FIXED: Use the new method to find and fill summary
            self.find_and_fill_summary_section(doc, summary)
            self.find_and_fill_innovativeness_section(doc, innovativeness)
            self.find_and_fill_keywords_section(doc, keywords)
            self.find_and_fill_ipr_section(doc, ipr)
            self.find_and_fill_commercialization_section(doc, commercialization)
            self.find_and_fill_literature_review_section(doc, literature_review)
            self.find_and_fill_collaboration_section(doc, collaboration)
            self.find_and_fill_literature_sources_section(doc, literature_sources)

            # Fill various sections and tables
            # ========================================================================
            # 5D. FILL TABLES AND SECTIONS
            # ========================================================================
            # WRITE YOUR TABLE FILLING LOGIC BELOW:
            for table in doc.tables:
                rows_list = list(table.rows)
                
                for row_idx, row in enumerate(rows_list):
                    row_text = ' '.join([cell.text for cell in row.cells]).strip()
                    
                    # Section 2.1.1 - Shareholder information
                    if "Shareholder" in row_text and "Company code" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 3:
                                next_row.cells[0].text = A_S_Ns_2
                                next_row.cells[1].text = company_code_2
                                next_row.cells[2].text = share_Hs_2

                    # Section 2.1.2 - Legal entity information
                    if "Name of the legal entity" in row_text and "Identification code" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 3:
                                next_row.cells[0].text = Name_of_the_legal_entity_2
                                next_row.cells[1].text = Identification_code_2
                                next_row.cells[2].text = Shareholding_Sh
                    
                    # Section 2.1.3 - Shareholder entities
                    if "2.1.3. Information on the legal entities in which the applicant's shareholders hold shares:" in row_text:
                        if "Name of the legal entity" in row_text and "Identification code" in row_text:
                            if row_idx + 1 < len(rows_list):
                                next_row = rows_list[row_idx + 1]
                                if len(next_row.cells) >= 3:
                                    next_row.cells[0].text = S_Hs_2
                                    next_row.cells[1].text = S_Is_2
                                    next_row.cells[2].text = S_Ss_2

                    # Section 2.2 - Current activities
                    elif "Activity(ies) carried out" in row_text and "Share of activity" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 3:
                                next_row.cells[0].text = main_activity_2
                                next_row.cells[1].text = activity_percentage_2 + "%"
                                next_row.cells[2].text = cese_class_2
                    
                    # Section 2.3 - Products offered
                    elif "Products offered" in row_text and "Percentage of sales" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 3:
                                next_row.cells[0].text = Pro_off_2
                                next_row.cells[1].text = Per_sales_2 + "%"
                                next_row.cells[2].text = "Lithuania, EU"
                    
                    # Section 3.1.2 - Level of novelty
                    elif "Product" in row_text and "Level of novelty" in row_text and "Justification" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 3:
                                next_row.cells[0].text = product_name_2
                                next_row.cells[1].text = novelty_level_2
                                next_row.cells[2].text = jus_pro_2

                    # Section 3.1.3 - R&D priority
                    elif "R&D priority" in row_text and "Justification" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 2:
                                next_row.cells[0].text = rd_priority_2
                                next_row.cells[1].text = jus_r_d_i_2
                    
                    # Section 4.2 - Existing staff
                    elif "Responsibilities" in row_text and "R&D activities" in row_text and "qualification" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 3:
                                next_row.cells[0].text = E_s_RES_2
                                next_row.cells[1].text = E_s_R_D_2
                                next_row.cells[2].text = E_s_R_2

                    # Section 4.2.2 - Staff with period
                    elif "Responsibilities" in row_text and "R&D activities" in row_text and "qualification" in row_text and "Period" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 4:
                                next_row.cells[0].text = A_s_RES_2
                                next_row.cells[1].text = A_s_R_D_2
                                next_row.cells[2].text = A_s_R_2
                                next_row.cells[3].text = A_s_P_2

                    # Section 4.2.3 - R&D personnel tasks
                    elif "Name of employee" in row_text and "Task(s)" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 6:
                                next_row.cells[0].text = "1"
                                next_row.cells[1].text = N_e_2
                                next_row.cells[2].text = N_r_2
                                next_row.cells[3].text = N_t_2
                                next_row.cells[4].text = N_w_T_2
                                next_row.cells[5].text = N_p_T_2
                    
                    # Section 5.2 - Competitors
                    elif "Name of competitor" in row_text and "Market share" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 3:
                                next_row.cells[0].text = competitor_2
                                next_row.cells[1].text = competitor_market_share_2 + "%"
                                next_row.cells[2].text = "Strong market position, higher prices, limited innovation"
                    
                    # Section 5.4 - Technological readiness levels
                    elif "Product" in row_text and "technological readiness" in row_text and "before" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 4:
                                next_row.cells[0].text = product_name_2
                                next_row.cells[1].text = current_tpl_2
                                next_row.cells[2].text = target_tpl_2   
                                next_row.cells[3].text = "Laboratory prototype validated, ready for pilot testing"
                    
                    # Section 6.1 - Main assets
                    elif "Name of asset" in row_text and "Form of ownership" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 4:
                                next_row.cells[0].text = N_As_2
                                next_row.cells[1].text = F_Os_2
                                next_row.cells[2].text = S_Us_2
                                next_row.cells[3].text = W_R_Ds_2

                    # Section 7.1 - Financial projections
                    elif "Planned revenue" in row_text and "Eligible project costs" in row_text:
                        if row_idx + 1 < len(rows_list):
                            next_row = rows_list[row_idx + 1]
                            if len(next_row.cells) >= 3:
                                next_row.cells[0].text = revenue_projection_2
                                next_row.cells[1].text = rd_budget_2
                                next_row.cells[2].text = revenue_ratio_2

                    # Section 4.5 - Project Impact table
                    elif "Project impact number(s) and title(s) specified in the PIP" in row_text:
                        # Set project impact title in the same row, second cell
                        if len(rows_list[row_idx].cells) > 1:
                            rows_list[row_idx].cells[1].text = project_impact_title_2

                        # Set start and completion month in the next row, cells 1 and 3
                        if row_idx + 1 < len(rows_list) and len(rows_list[row_idx + 1].cells) >= 4:
                            rows_list[row_idx + 1].cells[1].text = project_start_month_2
                            rows_list[row_idx + 1].cells[3].text = project_completion_month_2

                        # Set current and target TRL in the row after that, cells 1 and 3
                        if row_idx + 2 < len(rows_list) and len(rows_list[row_idx + 2].cells) >= 4:
                            rows_list[row_idx + 2].cells[1].text = current_tpl_2
                            rows_list[row_idx + 2].cells[3].text = target_tpl_2

                    # TPL Justification table
                    elif "Justification of the technological readiness level" in row_text:
                        if row_idx + 1 < len(rows_list) and len(rows_list[row_idx + 1].cells) > 1:
                            rows_list[row_idx + 1].cells[1].text = tpl_justification
                    
                    # Fill Risk assessment table (4.8) - with dynamic data
                    elif "Stages" in row_text and "Risks" in row_text and "Critical points" in row_text:
                        for i, (stage, risk, critical, mitigation) in enumerate(risk_stages):
                            if row_idx + 1 + i < len(rows_list):
                                risk_row = rows_list[row_idx + 1 + i]
                                if len(risk_row.cells) >= 4:
                                    risk_row.cells[0].text = stage
                                    risk_row.cells[1].text = risk
                                    risk_row.cells[2].text = critical
                                    risk_row.cells[3].text = mitigation
            
            # Save to BytesIO
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return {
                "success": True,
                "filled_document": doc_bytes,
                "filename": f"filled_mtep_business_plan_{company_name.replace(' ', '_')}.docx",
                "preview": f"MTEP Business Plan filled for {product_name} ({current_tpl} → {target_tpl})"
            }
            
        except Exception as e:
            print(f"MTEP filling error details: {e}")
            print(f"Error type: {type(e).__name__}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")
            return {
                "success": False,
                "error": f"Error filling MTEP business plan: {str(e)}"
            }
        finally:
            try:
                os.unlink(tmp_path)
            except:
                pass
    
    def fill_rd_assessment_form(self, uploaded_file, comprehensive_data):
        """Fill Document 3: R&D Assessment Form"""
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        try:
            doc = Document(tmp_path)
            
            # Extract data from comprehensive dataset
            project_type = comprehensive_data.get('PROJECT_TYPE', 'Information and communication technologies')
            project_subtopic = comprehensive_data.get('PROJECT_SUBTOPIC', 'Artificial intelligence, big and distributed data')
            rd_expenditure_2022 = comprehensive_data.get('RD_EXPENDITURE_2022', '200000')
            rd_expenditure_2023 = comprehensive_data.get('RD_EXPENDITURE_2023', '250000')
            total_research_jobs = comprehensive_data.get('TOTAL_RESEARCH_JOBS', '6')
            jobs_during_project = comprehensive_data.get('JOBS_DURING_PROJECT', '4')
            jobs_after_project = comprehensive_data.get('JOBS_AFTER_PROJECT', '2')
            manager_title = comprehensive_data.get('MANAGER_TITLE', 'R&D Director')
            manager_name = comprehensive_data.get('MANAGER_NAME', 'Dr. Vytautas Kazlauskas')
            
            # Define priority mappings
            priority_mapping = {
                "Health technologies and biotechnologies": {
                    "main_checkbox": "1.1. Health technologies and biotechnologies",
                    "subtopics": {
                        "Molecular technologies for medicine and biopharmaceuticals": "Molecular technologies for medicine and biopharmaceuticals",
                        "Advanced applied technologies for personal and public health": "1.1.2. Advanced applied technologies for personal and public health",
                        "Advanced medical engineering for early diagnosis and treatment": "1.1.3. Advanced medical engineering for early diagnosis and treatment",
                        "Safe food and sustainable agro-biological resources": "1.1.4. Safe food and sustainable agro-biological resources"
                    }
                },
                "Information and communication technologies": {
                    "main_checkbox": "1.3. Information and communication technologies",
                    "subtopics": {
                        "Artificial intelligence, big and distributed data": "1.3.1. Artificial intelligence, big and distributed data, heterogeneous analysis, processing and deployment",
                        "Internet of Things": "1.3.2. Internet of Things",
                        "Cyber security": "1.3.3. Cyber security",
                        "Financial technologies and blockchains": "1.3.4. Financial technologies and blockchains",
                        "Audiovisual media technologies and social innovation": "1.3.5. Audiovisual media technologies and social innovation",
                        "Intelligent transport systems": "1.3.6. Intelligent transport systems"
                    }
                }
            }
            
            # Fill tables with data
            for table in doc.tables:
                rows_list = list(table.rows)
                
                # Handle R&D priority selection table (Table 1)
                for row_idx, row in enumerate(rows_list):
                    row_text = ' '.join([cell.text for cell in row.cells]).strip()
                    
                    # Mark the appropriate main priority
                    if any(priority in row_text for priority in priority_mapping.keys()):
                        if project_type in row_text:
                            for cell in row.cells:
                                if '□' in cell.text and len(cell.text.strip()) < 5:
                                    cell.text = '☑'
                                    break
                    
                    # Mark the appropriate subtopic
                    for priority_key, priority_data in priority_mapping.items():
                        if project_type == priority_key:
                            for subtopic_key, subtopic_text in priority_data['subtopics'].items():
                                if subtopic_key.lower() in project_subtopic.lower() and subtopic_text in row_text:
                                    for cell in row.cells:
                                        if '□' in cell.text and len(cell.text.strip()) < 5:
                                            cell.text = '☑'
                                            break
                
                # Handle R&D expenditure table (Table 2)
                for row_idx, row in enumerate(rows_list):
                    row_text = ' '.join([cell.text for cell in row.cells]).strip()
                    
                    if "For 2022" in row_text and "For 2023" in row_text:
                        if row_idx + 1 < len(rows_list):
                            data_row = rows_list[row_idx + 1]
                            cells = list(data_row.cells)
                            if len(cells) >= 2:
                                if cells[0].text.strip() == "":
                                    cells[0].text = str(rd_expenditure_2022)
                                if cells[1].text.strip() == "":
                                    cells[1].text = str(rd_expenditure_2023)
                
                # Handle research jobs table (Table 3)
                for row_idx, row in enumerate(rows_list):
                    row_text = ' '.join([cell.text for cell in row.cells]).strip()
                    
                    if "Total number of research jobs created" in row_text and "Research jobs to be created in the project" in row_text:
                        if row_idx + 1 < len(rows_list):
                            data_row = rows_list[row_idx + 1]
                            cells = list(data_row.cells)
                            if len(cells) >= 3:
                                if cells[0].text.strip() == "":
                                    cells[0].text = str(total_research_jobs)
                                if cells[1].text.strip() == "":
                                    cells[1].text = str(jobs_during_project)
                                if cells[2].text.strip() == "":
                                    cells[2].text = str(jobs_after_project)
            
            # Fill signature area at the bottom
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if "title of manager" in text.lower() and "signature" in text.lower() and "name and surname" in text.lower():
                    new_text = text.replace("(title of manager or person authorised by him/her)", str(manager_title))
                    new_text = new_text.replace("(name and surname)", str(manager_name))
                    new_text = new_text.replace("(signature)", "[Signature]")
                    paragraph.text = new_text
            
            # Save to BytesIO
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return {
                "success": True,
                "filled_document": doc_bytes,
                "filename": f"filled_rd_assessment_{project_type.replace(' ', '_')}.docx",
                "preview": f"R&D Assessment: {project_subtopic}, Jobs: {total_research_jobs}, Budget: €{rd_expenditure_2023}"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error filling R&D assessment: {str(e)}"
            }
        finally:
            try:
                os.unlink(tmp_path)
            except:
                pass
    
    def process_passthrough_document(self, uploaded_file):
        """Process Document 4: Pass-through (no filling, just upload/download)"""
        try:
            # Just return the uploaded file as-is
            file_bytes = BytesIO()
            file_bytes.write(uploaded_file.getvalue())
            file_bytes.seek(0)
            
            # Get original filename or create a default one
            original_filename = uploaded_file.name if hasattr(uploaded_file, 'name') else "document_4.docx"
            download_filename = f"processed_{original_filename}"
            
            return {
                "success": True,
                "filled_document": file_bytes,
                "filename": download_filename,
                "preview": f"Document uploaded and ready for download: {original_filename}"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error processing document: {str(e)}"
            }
    
    def process_all_documents(self, declaration_file, mtep_file, rd_assessment_file, passthrough_file, comprehensive_data):
        """Process all four documents (3 with filling + 1 pass-through)"""
        results = {
            "declaration_result": None,
            "mtep_result": None,
            "rd_assessment_result": None,
            "passthrough_result": None
        }
        
        # Fill declaration document
        if declaration_file:
            results["declaration_result"] = self.fill_declaration_form(declaration_file, comprehensive_data)
        
        # Fill MTEP business plan
        if mtep_file:
            results["mtep_result"] = self.fill_mtep_business_plan(mtep_file, comprehensive_data)
        
        # Fill R&D assessment
        if rd_assessment_file:
            results["rd_assessment_result"] = self.fill_rd_assessment_form(rd_assessment_file, comprehensive_data)
        
        # Process pass-through document (no filling)
        if passthrough_file:
            results["passthrough_result"] = self.process_passthrough_document(passthrough_file)
        
        return results



   
    def process_all_documents_from_paths(self, declaration_path, mtep_path, rd_assessment_path, passthrough_path, comprehensive_data):
        """Process all four documents from file paths (3 with filling + 1 pass-through)"""
        results = {
            "declaration_result": None,
            "mtep_result": None,
            "rd_assessment_result": None,
            "passthrough_result": None
        }
        
        # Helper function to convert file path to file-like object
        def path_to_file_like(file_path):
            if os.path.exists(file_path):
                with open(file_path, 'rb') as f:
                    content = f.read()
                return BytesIO(content)
            return None
        
        # Fill declaration document
        if declaration_path and os.path.exists(declaration_path):
            declaration_file_like = path_to_file_like(declaration_path)
            if declaration_file_like:
                results["declaration_result"] = self.fill_declaration_form(declaration_file_like, comprehensive_data)
        
        # Fill MTEP business plan
        if mtep_path and os.path.exists(mtep_path):
            mtep_file_like = path_to_file_like(mtep_path)
            if mtep_file_like:
                results["mtep_result"] = self.fill_mtep_business_plan(mtep_file_like, comprehensive_data)
        
        # Fill R&D assessment
        if rd_assessment_path and os.path.exists(rd_assessment_path):
            rd_assessment_file_like = path_to_file_like(rd_assessment_path)
            if rd_assessment_file_like:
                results["rd_assessment_result"] = self.fill_rd_assessment_form(rd_assessment_file_like, comprehensive_data)
        
        # Pass-through document (no filling)
        if passthrough_path and os.path.exists(passthrough_path):
            passthrough_file_like = path_to_file_like(passthrough_path)
            if passthrough_file_like:
                results["passthrough_result"] = self.process_passthrough_document(passthrough_file_like)
        
        return results






# # ============================================================================
# # COMPREHENSIVE DOCUMENT FILLING AGENT - STRUCTURED FORMAT
# # ============================================================================

# from utils.ai_helper import AIHelper
# from docx import Document
# import tempfile
# import os
# from datetime import datetime
# import re
# from io import BytesIO

# class ComprehensiveDocumentFillingAgent:
    
#     # ============================================================================
#     # 1. INITIALIZATION
#     # ============================================================================
#     def __init__(self):
#         # WRITE YOUR INITIALIZATION CODE BELOW:
#         self.ai_helper = AIHelper()
    
#     # ============================================================================
#     # 2. DECLARATION FORM FILLING (DOCUMENT 1)
#     # ============================================================================
#     def fill_declaration_form(self, uploaded_file, comprehensive_data):
#         """Fill Document 1: Declaration Form"""
        
#         # WRITE YOUR DECLARATION FILLING LOGIC BELOW:
#         with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
#             tmp_file.write(uploaded_file.getvalue())
#             tmp_path = tmp_file.name
        
#         try:
#             doc = Document(tmp_path)
            
#             # Extract data from comprehensive dataset
#             company_name = comprehensive_data.get('COMPANY_NAME', 'Baltic Innovation Technologies UAB')
#             company_code = comprehensive_data.get('COMPANY_CODE', 'LT305123456')
#             completion_date = comprehensive_data.get('COMPLETION_DATE', datetime.now().strftime("%d.%m.%Y"))
#             manager_name = comprehensive_data.get('MANAGER_NAME', 'Dr. Vytautas Kazlauskas')
#             manager_position = comprehensive_data.get('MANAGER_POSITION', 'CEO')
            
#             # Fill the date at the top
#             date_filled = False
#             for paragraph in doc.paragraphs:
#                 text = paragraph.text
#                 if "date of completion" in text.lower():
#                     new_text = re.sub(r'_{3,}', completion_date, text)
#                     if new_text != text:
#                         paragraph.text = new_text
#                         date_filled = True
#                         break
            
#             # If not found, look for any paragraph with multiple underscores near the top
#             if not date_filled:
#                 for i, paragraph in enumerate(doc.paragraphs[:10]):
#                     text = paragraph.text
#                     if re.search(r'_{10,}', text) and len(text.strip()) < 100:
#                         new_text = re.sub(r'_{10,}', completion_date, text)
#                         paragraph.text = new_text
#                         date_filled = True
#                         break
            
#             # Fill the table with company information
#             for table in doc.tables:
#                 rows_list = list(table.rows)
                
#                 for row_idx, row in enumerate(rows_list):
#                     row_text = ' '.join([cell.text for cell in row.cells]).strip()
                    
#                     # Company name field
#                     if "Name of the declaring undertaking" in row_text:
#                         if row_idx + 1 < len(rows_list):
#                             next_row = rows_list[row_idx + 1]
#                             for cell in next_row.cells:
#                                 if cell.text.strip() == "" or len(cell.text.strip()) < 5:
#                                     cell.text = "      " + company_name
#                                     break
                    
#                     # Company code field
#                     elif "Code of the declaring company" in row_text:
#                         if row_idx + 1 < len(rows_list):
#                             next_row = rows_list[row_idx + 1]
#                             for cell in next_row.cells:
#                                 if cell.text.strip() == "" or len(cell.text.strip()) < 5:
#                                     cell.text = "      " + company_code
#                                     break
            
#             # Fill signature area
#             for paragraph in doc.paragraphs:
#                 text = paragraph.text
#                 if "position of the manager" in text.lower() or ("signature" in text.lower() and "_" in text):
#                     signature_line = f"{manager_position}                                    {manager_name}"
#                     new_text = re.sub(r'_{20,}.*?_{10,}', signature_line, text)
#                     if new_text != text:
#                         paragraph.text = new_text
            
#             # Save to BytesIO
#             doc_bytes = BytesIO()
#             doc.save(doc_bytes)
#             doc_bytes.seek(0)
            
#             return {
#                 "success": True,
#                 "filled_document": doc_bytes,
#                 "filename": f"filled_declaration_{company_name.replace(' ', '_')}.docx",
#                 "preview": f"Declaration filled for {company_name} (Code: {company_code})"
#             }
            
#         except Exception as e:
#             return {
#                 "success": False,
#                 "error": f"Error filling declaration document: {str(e)}"
#             }
#         finally:
#             try:
#                 os.unlink(tmp_path)
#             except:
#                 pass
    
#     # ============================================================================
#     # 3. RESEARCH AREAS HELPER FUNCTION
#     # ============================================================================
#     def fill_research_areas_in_document(self, doc, comprehensive_data):
#         """Fill research areas in the MTEP document"""
        
#         # WRITE YOUR RESEARCH AREAS LOGIC BELOW:
#         research_area = comprehensive_data.get('RESEARCH_AREA', 'Computer Sciences')
        
#         # Look for research area table
#         for table in doc.tables:
#             for row_idx, row in enumerate(table.rows):
#                 row_text = ' '.join([cell.text for cell in row.cells]).strip()
                
#                 # Find the research areas table header
#                 if "Research area(s) of the project" in row_text and "Indicate up to two areas" in row_text:
#                     if row_idx + 1 < len(table.rows):
#                         content_row = table.rows[row_idx + 1]
#                         if len(content_row.cells) >= 1:
#                             content_row.cells[0].text = f"1. {research_area}(main)\n2. Artificial Intelligence"
#                             break
                
#                 # Look for cells that contain research area patterns
#                 for cell in row.cells:
#                     if "Research area(s) of the project" in cell.text:
#                         cell_text = cell.text.strip()
#                         if "1." not in cell_text:
#                             cell_idx = list(row.cells).index(cell)
#                             if cell_idx + 1 < len(row.cells):
#                                 row.cells[cell_idx + 1].text = f"1. {research_area}(main)\n2. Artificial Intelligence"
#                                 return
#                             if row_idx + 1 < len(table.rows):
#                                 next_row = table.rows[row_idx + 1]
#                                 if len(next_row.cells) >= 1:
#                                     next_row.cells[0].text = f"1. {research_area}(main)\n2. Artificial Intelligence"
#                                     return

#     # ============================================================================
#     # 4. PARAGRAPH TEXT HELPER FUNCTIONS
#     # ============================================================================
#     def add_text_to_paragraph(self, paragraph, text_to_add):
#         """Helper method to add text to an existing paragraph"""
        
#         # WRITE YOUR PARAGRAPH TEXT LOGIC BELOW:
#         if paragraph.text.strip():
#             paragraph.text = paragraph.text + "\n\n" + text_to_add
#         else:
#             paragraph.text = text_to_add

#     def find_and_fill_summary_section(self, doc, summary):
#         """Find and fill the summary section in the document"""
        
#         # WRITE YOUR SUMMARY FILLING LOGIC BELOW:
#         summary_filled = False
        
#         # Method 1: Look for specific summary patterns in paragraphs
#         for i, para in enumerate(doc.paragraphs):
#             para_text = para.text.lower().strip()
            
#             # Look for summary-related text patterns
#             if any(pattern in para_text for pattern in [
#                 "what are the long-term objectives set",
#                 "long-term objectives",
#                 "project summary",
#                 "summary of the project",
#                 "describe the project"
#             ]):
#                 # Try to find the next empty or nearly empty paragraph
#                 for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
#                     next_para = doc.paragraphs[j]
#                     if len(next_para.text.strip()) < 20:  # Empty or very short paragraph
#                         next_para.text = summary
#                         summary_filled = True
#                         break
                
#                 if summary_filled:
#                     break
                
#                 # If no empty paragraph found, add to current paragraph
#                 self.add_text_to_paragraph(para, f"\n\n{summary}")
#                 summary_filled = True
#                 break
        
#         # Method 2: If not found in paragraphs, look in table cells
#         if not summary_filled:
#             for table in doc.tables:
#                 for row in table.rows:
#                     for cell in row.cells:
#                         cell_text = cell.text.lower().strip()
#                         if any(pattern in cell_text for pattern in [
#                             "objectives",
#                             "summary",
#                             "description",
#                             "project goals"
#                         ]):
#                             if len(cell.text.strip()) < 50:  # Relatively empty cell
#                                 cell.text = summary
#                                 summary_filled = True
#                                 break
#                     if summary_filled:
#                         break
#                 if summary_filled:
#                     break
        
#         # Method 3: If still not found, add to first suitable paragraph
#         if not summary_filled:
#             for i, para in enumerate(doc.paragraphs):
#                 if len(para.text.strip()) < 10 and i > 5:  # Skip header paragraphs
#                     para.text = f"Project Summary:\n{summary}"
#                     summary_filled = True
#                     break
        
#         return summary_filled


#     def find_and_fill_innovativeness_section(self, doc, innovativeness):
#         """Find and fill the innovativeness section in the document (Section 3.1.1)"""
        
#         # WRITE YOUR INNOVATIVENESS FILLING LOGIC BELOW:
#         innovativeness_filled = False
        
#         # Method 1: Find the exact text "new products will bring to consumers, what problems they will solve, etc." and fill below it
#         for i, para in enumerate(doc.paragraphs):
#             para_text = para.text.lower().strip()
            
#             # Look for the specific ending text pattern
#             if "new products will bring to consumers, what problems they will solve" in para_text:
#                 # Found the target paragraph, now fill the next empty paragraph or create content below
                
#                 # Try to find the next empty or nearly empty paragraph
#                 content_filled = False
#                 for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
#                     next_para = doc.paragraphs[j]
#                     if len(next_para.text.strip()) < 30:  # Empty or very short paragraph
#                         next_para.text = innovativeness
#                         innovativeness_filled = True
#                         content_filled = True
#                         break
                
#                 # If no empty paragraph found, add content directly to the current paragraph
#                 if not content_filled:
#                     if para.text.strip().endswith('.') or para.text.strip().endswith('etc.'):
#                         para.text = para.text + "\n\n" + innovativeness
#                     else:
#                         para.text = para.text + " " + innovativeness
#                     innovativeness_filled = True
#                 break
            
#         return innovativeness_filled

#     def find_and_fill_keywords_section(self, doc, keywords):
#         """Find and fill the keywords section in the document (Section 3.1.5)"""
        
#         # WRITE YOUR KEYWORDS FILLING LOGIC BELOW:
#         keywords_filled = False
        
#         # Method 1: Look for specific keywords patterns in paragraphs
#         for i, para in enumerate(doc.paragraphs):
#             para_text = para.text.lower().strip()
            
#             # Look for keywords-related text patterns
#             if "keywords" in para_text or "key words" in para_text:
#                 # Try to find the next empty or nearly empty paragraph
#                 for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
#                     next_para = doc.paragraphs[j]
#                     if len(next_para.text.strip()) < 20:
#                         next_para.text = keywords
#                         keywords_filled = True
#                         break
                
#                 if not keywords_filled:
#                     if para.text.strip().endswith('.'):
#                         para.text = para.text + "\n\n" + keywords
#                     else:
#                         para.text = para.text + " " + keywords
#                     keywords_filled = True
#                 break
#         return keywords_filled
#     # ============================================================================
#     # 5. MTEP BUSINESS PLAN FILLING (DOCUMENT 2) - MAIN FUNCTION
#     # ============================================================================
#     def fill_mtep_business_plan(self, uploaded_file, comprehensive_data):
#         """Fill Document 2: MTEP Business Plan"""
        
#         # WRITE YOUR MTEP FILLING LOGIC BELOW:
#         with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
#             tmp_file.write(uploaded_file.getvalue())
#             tmp_path = tmp_file.name
        
#         try:
#             doc = Document(tmp_path)
            
#             # Fill research areas first
#             self.fill_research_areas_in_document(doc, comprehensive_data)
        
#             # ========================================================================
#             # 5A. EXTRACT DATA FROM COMPREHENSIVE DATASET
#             # ========================================================================
#             # WRITE YOUR DATA EXTRACTION BELOW:
#             company_name = str(comprehensive_data.get('COMPANY_NAME', 'Baltic Innovation Technologies UAB'))
#             company_code_2 = str(comprehensive_data.get('COMPANY_CODE', 'LT305123456'))
#             S_Ss_2 = str(comprehensive_data.get('S_S', '50'))
#             S_Is_2 = str(comprehensive_data.get('S_I', 'MT23456'))
#             S_Hs_2 = str(comprehensive_data.get('S_H', '50'))
#             Name_of_the_legal_entity_2 = str(comprehensive_data.get('N_L_E', company_name))
#             Identification_code_2 = str(comprehensive_data.get('I_C', 'LT123456789'))
#             Shareholding_Sh = str(comprehensive_data.get('Sharehol', '50'))
#             A_S_Ns_2 = str(comprehensive_data.get('A_S_Ns', 'Vytautas Kazlauskas'))
#             share_Hs_2 = str(comprehensive_data.get('SHARE_HS', '100'))
#             summary_2 = str(comprehensive_data.get('SUMMARY_1', 'AAAAAAAAAAABBBBBBBBBBBCCCCCCCCCCC'))
#             innovativeness_2 = str(comprehensive_data.get('INNOVATIVENESS', 'Our product features advanced AI algorithms with real-time processing capabilities, offering 95% accuracy compared to 70% of existing solutions. The innovative machine learning architecture provides personalized recommendations, solving the problem of generic one-size-fits-all approaches. This brings significant benefits to consumers through improved efficiency, cost reduction of up to 40%, and enhanced user experience with intuitive interfaces.'))
#             keywords = str(comprehensive_data.get('KEYWORDS', 'AI, machine learning, real-time processing, personalized recommendations, efficiency, cost reduction'))
#             main_activity_2 = str(comprehensive_data.get('MAIN_ACTIVITY', 'Software development and innovation'))
#             activity_percentage_2 = str(comprehensive_data.get('ACTIVITY_PERCENTAGE', '75'))
#             cese_class_2 = str(comprehensive_data.get('CESE_CLASS', '62.01'))
#             Pro_off_2 = str(comprehensive_data.get('PRODUCTS_OFFERED', 'AI-powered diagnostic system'))
#             Per_sales_2 = str(comprehensive_data.get('PER_SALES', '100'))
#             product_name_2 = str(comprehensive_data.get('PRODUCT_NAME', 'AI-powered diagnostic system'))
#             novelty_level_2 = str(comprehensive_data.get('NOVELTY_LEVEL', 'market level'))
#             jus_pro_2 = str(comprehensive_data.get('JUS_PRO', 'This product introduces new technological solutions not available in the current market'))
#             rd_priority_2 = str(comprehensive_data.get('RD_PRIORITY', 'Information and communication technologies'))
#             jus_r_d_i_2 = str(comprehensive_data.get('JUS_R_D_I', 'This project aligns with smart specialization priorities in ICT and digital innovation'))
#             N_As_2 = str(comprehensive_data.get('N_As', 'R&D Laboratory'))
#             F_Os_2 = str(comprehensive_data.get('F_O', 'Owned'))
#             S_Us_2 = str(comprehensive_data.get('S_U', '200 m²'))
#             W_R_Ds_2 = str(comprehensive_data.get('W_R_D', 'All R&D activities, testing, prototype development'))            
#             competitor_2 = str(comprehensive_data.get('COMPETITOR_M', 'TechCorp Europe'))
#             competitor_market_share_2 = str(comprehensive_data.get('COMPETITOR_MARKET_SHARE', '25'))
#             current_tpl_2 = str(comprehensive_data.get('CURRENT_TPL', 'TPL 3'))
#             target_tpl_2 = str(comprehensive_data.get('TARGET_TPL', 'TPL 6'))
#             revenue_projection_2 = str(comprehensive_data.get('REVENUE_PROJECTION', '500000'))
#             rd_budget_2 = str(comprehensive_data.get('RD_BUDGET', '200000'))
#             revenue_ratio_2 = str(comprehensive_data.get('REVENUE_RATIO', '2.5'))
#             project_impact_title_2 = str(comprehensive_data.get('PROJECT_IMPACT_TITLE', f'Development of {product_name_2}'))
#             project_start_month_2 = str(comprehensive_data.get('PROJECT_START_MONTH', '1'))
#             project_completion_month_2 = str(comprehensive_data.get('PROJECT_COMPLETION_MONTH', '24'))
#             E_s_RES_2 = str(comprehensive_data.get('E_S_RES', 'NOT_FOUND'))
#             E_s_R_D_2 = str(comprehensive_data.get('E_S_R&D', 'NOT_FOUND'))
#             E_s_R_2 = str(comprehensive_data.get('E_S_R', 'NOT_FOUND'))
#             A_s_RES_2 = str(comprehensive_data.get('A_S_RES', 'NOT_FOUND'))
#             A_s_R_D_2 = str(comprehensive_data.get('A_S_R&D', 'NOT_FOUND'))
#             A_s_R_2 = str(comprehensive_data.get('A_S_R', 'NOT_FOUND'))
#             A_s_P_2 = str(comprehensive_data.get('A_S_P', 'NOT_FOUND'))

#             N_E_2 = str(comprehensive_data.get('N_E', 'NOT_FOUND'))
#             N_R_2 = str(comprehensive_data.get('N_R', 'NOT_FOUND'))
#             N_T_2 = str(comprehensive_data.get('N_T', 'NOT_FOUND'))
#             N_W_T_2 = str(comprehensive_data.get('N_W_T', 'NOT_FOUND'))
#             N_P_T_2 = str(comprehensive_data.get('N_P_T', 'NOT_FOUND'))



#             project_impact_description = str(comprehensive_data.get('PROJECT_IMPACT_DESCRIPTION', 
#                 f'The objective is to develop {product_name_2} through systematic R&D activities.'))
#             tpl_justification = str(comprehensive_data.get('TPL_JUSTIFICATION', 
#                 f'Current {current_tpl_2} achieved through laboratory validation. Target {target_tpl_2} through systematic development.'))
            
#             # Clean percentage values
#             activity_percentage_2 = activity_percentage_2.replace('%', '').strip()
#             competitor_market_share_2 = competitor_market_share_2.replace('%', '').strip()

#             # ========================================================================
#             # 5B. RISK ASSESSMENT DATA EXTRACTION
#             # ========================================================================
#             # WRITE YOUR RISK ASSESSMENT DATA BELOW:
#             risk_stages = [
#                 (
#                     str(comprehensive_data.get('RISK_STAGE_1', 'Concept formulation and feasibility validation')),
#                     str(comprehensive_data.get('RISK_DESCRIPTION_1', 'Market acceptance uncertainty')),
#                     str(comprehensive_data.get('CRITICAL_POINT_1', 'User needs validation')),
#                     str(comprehensive_data.get('MITIGATION_ACTION_1', 'Extensive market research'))
#                 ),
#                 (
#                     str(comprehensive_data.get('RISK_STAGE_2', 'Layout development, testing, checking')),
#                     str(comprehensive_data.get('RISK_DESCRIPTION_2', 'Technical complexity')),
#                     str(comprehensive_data.get('CRITICAL_POINT_2', 'Algorithm optimization')),
#                     str(comprehensive_data.get('MITIGATION_ACTION_2', 'Iterative testing approach'))
#                 ),
#                 (
#                     str(comprehensive_data.get('RISK_STAGE_3', 'Prototype development and demonstration')),
#                     str(comprehensive_data.get('RISK_DESCRIPTION_3', 'Integration challenges')),
#                     str(comprehensive_data.get('CRITICAL_POINT_3', 'System compatibility')),
#                     str(comprehensive_data.get('MITIGATION_ACTION_3', 'Modular development'))
#                 ),
#                 (
#                     str(comprehensive_data.get('RISK_STAGE_4', 'Production and evaluation of a pilot batch')),
#                     str(comprehensive_data.get('RISK_DESCRIPTION_4', 'Scale-up difficulties')),
#                     str(comprehensive_data.get('CRITICAL_POINT_4', 'Performance at scale')),
#                     str(comprehensive_data.get('', 'Pilot testing program'))
#                 )
#             ]

#             # ==================<FILL SUMMARY SECTION=======================================
#             # 5C. FILL SUMMARY SECTION
#             # ========================================================================
#             # WRITE YOUR SUMMARY FILLING LOGIC BELOW:
#             self.find_and_fill_summary_section(doc, summary_2)
#             self.find_and_fill_innovativeness_section(doc, innovativeness_2)
#             self.find_and_fill_keywords_section(doc, keywords)
            # # ========================================================================
            # # 5D. FILL TABLES AND SECTIONS
            # # ========================================================================
            # # WRITE YOUR TABLE FILLING LOGIC BELOW:
            # for table in doc.tables:
            #     rows_list = list(table.rows)
                
            #     for row_idx, row in enumerate(rows_list):
            #         row_text = ' '.join([cell.text for cell in row.cells]).strip()
                    
            #         # Section 2.1.1 - Shareholder information
            #         if "Shareholder" in row_text and "Company code" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 3:
            #                     next_row.cells[0].text = A_S_Ns_2
            #                     next_row.cells[1].text = company_code_2
            #                     next_row.cells[2].text = share_Hs_2

            #         # Section 2.1.2 - Legal entity information
            #         if "Name of the legal entity" in row_text and "Identification code" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 3:
            #                     next_row.cells[0].text = Name_of_the_legal_entity_2
            #                     next_row.cells[1].text = Identification_code_2
            #                     next_row.cells[2].text = Shareholding_Sh
                    
            #         # Section 2.1.3 - Shareholder entities
            #         if "2.1.3. Information on the legal entities in which the applicant's shareholders hold shares:" in row_text:
            #             if "Name of the legal entity" in row_text and "Identification code" in row_text:
            #                 if row_idx + 1 < len(rows_list):
            #                     next_row = rows_list[row_idx + 1]
            #                     if len(next_row.cells) >= 3:
            #                         next_row.cells[0].text = S_Hs_2
            #                         next_row.cells[1].text = S_Is_2
            #                         next_row.cells[2].text = S_Ss_2

            #         # Section 2.2 - Current activities
            #         elif "Activity(ies) carried out" in row_text and "Share of activity" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 3:
            #                     next_row.cells[0].text = main_activity_2
            #                     next_row.cells[1].text = activity_percentage_2 + "%"
            #                     next_row.cells[2].text = cese_class_2
                    
            #         # Section 2.3 - Products offered
            #         elif "Products offered" in row_text and "Percentage of sales" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 3:
            #                     next_row.cells[0].text = Pro_off_2
            #                     next_row.cells[1].text = Per_sales_2 + "%"
            #                     next_row.cells[2].text = "Lithuania, EU"
                    
            #         # Section 3.1.2 - Level of novelty
            #         elif "Product" in row_text and "Level of novelty" in row_text and "Justification" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 3:
            #                     next_row.cells[0].text = product_name_2
            #                     next_row.cells[1].text = novelty_level_2
            #                     next_row.cells[2].text = jus_pro_2

            #         # Section 3.1.3 - R&D priority
            #         elif "R&D priority" in row_text and "Justification" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 2:
            #                     next_row.cells[0].text = rd_priority_2
            #                     next_row.cells[1].text = jus_r_d_i_2
                    
            #         # Section 4.2 - Existing staff
            #         elif "Responsibilities" in row_text and "R&D activities" in row_text and "qualification" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 3:
            #                     next_row.cells[0].text = E_s_RES_2
            #                     next_row.cells[1].text = E_s_R_D_2
            #                     next_row.cells[2].text = E_s_R_2

            #         # Section 4.2.2 - Staff with period
            #         elif "Responsibilities" in row_text and "R&D activities" in row_text and "qualification" in row_text and "Period" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 4:
            #                     next_row.cells[0].text = A_s_RES_2
            #                     next_row.cells[1].text = A_s_R_D_2
            #                     next_row.cells[2].text = A_s_R_2
            #                     next_row.cells[3].text = A_s_P_2

            #         # Section 4.2.3 - R&D personnel tasks
            #         elif "Name of employee" in row_text and "Task(s)" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 6:
            #                     next_row.cells[0].text = "1"
            #                     next_row.cells[1].text = N_E_2
            #                     next_row.cells[2].text = N_R_2
            #                     next_row.cells[3].text = N_T_2
            #                     next_row.cells[4].text = N_W_T_2
            #                     next_row.cells[5].text = N_P_T_2
                    
            #         # Section 5.2 - Competitors
            #         elif "Name of competitor" in row_text and "Market share" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 3:
            #                     next_row.cells[0].text = competitor_2
            #                     next_row.cells[1].text = competitor_market_share_2 + "%"
            #                     next_row.cells[2].text = "Strong market position, higher prices, limited innovation"
                    
            #         # Section 5.4 - Technological readiness levels
            #         elif "Product" in row_text and "technological readiness" in row_text and "before" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 4:
            #                     next_row.cells[0].text = product_name_2
            #                     next_row.cells[1].text = current_tpl_2
            #                     next_row.cells[2].text = target_tpl_2   
            #                     next_row.cells[3].text = "Laboratory prototype validated, ready for pilot testing"
                    
            #         # Section 6.1 - Main assets
            #         elif "Name of asset" in row_text and "Form of ownership" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 4:
            #                     next_row.cells[0].text = N_As_2
            #                     next_row.cells[1].text = F_Os_2
            #                     next_row.cells[2].text = S_Us_2
            #                     next_row.cells[3].text = W_R_Ds_2

            #         # Section 7.1 - Financial projections
            #         elif "Planned revenue" in row_text and "Eligible project costs" in row_text:
            #             if row_idx + 1 < len(rows_list):
            #                 next_row = rows_list[row_idx + 1]
            #                 if len(next_row.cells) >= 3:
            #                     next_row.cells[0].text = revenue_projection_2
            #                     next_row.cells[1].text = rd_budget_2
            #                     next_row.cells[2].text = revenue_ratio_2

            #         # Section 4.5 - Project Impact table
            #         elif "Project impact number(s) and title(s) specified in the PIP" in row_text:
            #             # Set project impact title in the same row, second cell
            #             if len(rows_list[row_idx].cells) > 1:
            #                 rows_list[row_idx].cells[1].text = project_impact_title_2

            #             # Set start and completion month in the next row, cells 1 and 3
            #             if row_idx + 1 < len(rows_list) and len(rows_list[row_idx + 1].cells) >= 4:
            #                 rows_list[row_idx + 1].cells[1].text = project_start_month_2
            #                 rows_list[row_idx + 1].cells[3].text = project_completion_month_2

            #             # Set current and target TRL in the row after that, cells 1 and 3
            #             if row_idx + 2 < len(rows_list) and len(rows_list[row_idx + 2].cells) >= 4:
            #                 rows_list[row_idx + 2].cells[1].text = current_tpl_2
            #                 rows_list[row_idx + 2].cells[3].text = target_tpl_2

            #         # TPL Justification table
            #         elif "Justification of the technological readiness level" in row_text:
            #             if row_idx + 1 < len(rows_list) and len(rows_list[row_idx + 1].cells) > 1:
            #                 rows_list[row_idx + 1].cells[1].text = tpl_justification
                    
            #         # Fill Risk assessment table (4.8) - with dynamic data
            #         elif "Stages" in row_text and "Risks" in row_text and "Critical points" in row_text:
            #             for i, (stage, risk, critical, mitigation) in enumerate(risk_stages):
            #                 if row_idx + 1 + i < len(rows_list):
            #                     risk_row = rows_list[row_idx + 1 + i]
            #                     if len(risk_row.cells) >= 4:
            #                         risk_row.cells[0].text = stage
            #                         risk_row.cells[1].text = risk
            #                         risk_row.cells[2].text = critical
            #                         risk_row.cells[3].text = mitigation
            
#             # ========================================================================
#             # 5E. SAVE AND RETURN
#             # ========================================================================
#             # WRITE YOUR SAVE LOGIC BELOW:
#             doc_bytes = BytesIO()
#             doc.save(doc_bytes)
#             doc_bytes.seek(0)
            
#             return {
#                 "success": True,
#                 "filled_document": doc_bytes,
#                 "filename": f"filled_mtep_business_plan_{company_name.replace(' ', '_')}.docx",
#                 "preview": f"MTEP Business Plan filled for {product_name_2} ({current_tpl_2} → {target_tpl_2})"
#             }
            
#         except Exception as e:
#             print(f"MTEP filling error details: {e}")
#             print(f"Error type: {type(e).__name__}")
#             import traceback
#             print(f"Traceback: {traceback.format_exc()}")
#             return {
#                 "success": False,
#                 "error": f"Error filling MTEP business plan: {str(e)}"
#             }
#         finally:
#             try:
#                 os.unlink(tmp_path)
#             except:
#                 pass
    
#     # ============================================================================
#     # 6. R&D ASSESSMENT FORM FILLING (DOCUMENT 3)
#     # ============================================================================
#     def fill_rd_assessment_form(self, uploaded_file, comprehensive_data):
#         """Fill Document 3: R&D Assessment Form"""
        
#         # WRITE YOUR R&D ASSESSMENT FILLING LOGIC BELOW:
#         with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
#             tmp_file.write(uploaded_file.getvalue())
#             tmp_path = tmp_file.name
        
#         try:
#             doc = Document(tmp_path)
            
#             # ========================================================================
#             # 6A. EXTRACT R&D ASSESSMENT DATA
#             # ========================================================================
#             # WRITE YOUR R&D DATA EXTRACTION BELOW:
#             project_type = comprehensive_data.get('PROJECT_TYPE', 'Information and communication technologies')
#             project_subtopic = comprehensive_data.get('PROJECT_SUBTOPIC', 'Artificial intelligence, big and distributed data')
#             rd_expenditure_2022 = comprehensive_data.get('RD_EXPENDITURE_2022', '200000')
#             rd_expenditure_2023 = comprehensive_data.get('RD_EXPENDITURE_2023', '250000')
#             total_research_jobs = comprehensive_data.get('TOTAL_RESEARCH_JOBS', '6')
#             jobs_during_project = comprehensive_data.get('JOBS_DURING_PROJECT', '4')
#             jobs_after_project = comprehensive_data.get('JOBS_AFTER_PROJECT', '2')
#             manager_title = comprehensive_data.get('MANAGER_TITLE', 'R&D Director')
#             manager_name = comprehensive_data.get('MANAGER_NAME', 'Dr. Vytautas Kazlauskas')
            
#             # ========================================================================
#             # 6B. PRIORITY MAPPINGS DEFINITION
#             # ========================================================================
#             # WRITE YOUR PRIORITY MAPPINGS BELOW:
#             priority_mapping = {
#                 "Health technologies and biotechnologies": {
#                     "main_checkbox": "1.1. Health technologies and biotechnologies",
#                     "subtopics": {
#                         "Molecular technologies for medicine and biopharmaceuticals": "Molecular technologies for medicine and biopharmaceuticals",
#                         "Advanced applied technologies for personal and public health": "1.1.2. Advanced applied technologies for personal and public health",
#                         "Advanced medical engineering for early diagnosis and treatment": "1.1.3. Advanced medical engineering for early diagnosis and treatment",
#                         "Safe food and sustainable agro-biological resources": "1.1.4. Safe food and sustainable agro-biological resources"
#                     }
#                 },
#                 "New production processes, materials and technologies": {
#                     "main_checkbox": "1.2. New production processes, materials and technologies",
#                     "subtopics": {
#                         "Photonics and laser technologies": "Photonics and laser technologies",
#                         "Advanced materials and structures": "1.2.2. Advanced materials and structures",
#                         "Flexible technologies for product development": "1.2.3. Flexible technologies for product development, production and process management, design",
#                         "Energy efficiency, smartness": "1.2.4. Energy efficiency, smartness",
#                         "Renewable energy sources": "1.2.5. Renewable energy sources"
#                     }
#                 },
#                 "Information and communication technologies": {
#                     "main_checkbox": "1.3. Information and communication technologies",
#                     "subtopics": {
#                         "Artificial intelligence, big and distributed data": "1.3.1. Artificial intelligence, big and distributed data, heterogeneous analysis, processing and deployment",
#                         "Internet of Things": "1.3.2. Internet of Things",
#                         "Cyber security": "1.3.3. Cyber security",
#                         "Financial technologies and blockchains": "1.3.4. Financial technologies and blockchains",
#                         "Audiovisual media technologies and social innovation": "1.3.5. Audiovisual media technologies and social innovation",
#                         "Intelligent transport systems": "1.3.6. Intelligent transport systems"
#                     }
#                 }
#             }
            
#             # ========================================================================
#             # 6C. FILL TABLES WITH DATA
#             # ========================================================================
#             # WRITE YOUR TABLE FILLING LOGIC BELOW:
#             for table in doc.tables:
#                 rows_list = list(table.rows)
                
#                 # Handle R&D priority selection table (Table 1)
#                 for row_idx, row in enumerate(rows_list):
#                     row_text = ' '.join([cell.text for cell in row.cells]).strip()
                    
#                     # Mark the appropriate main priority
#                     if any(priority in row_text for priority in priority_mapping.keys()):
#                         if project_type in row_text:
#                             for cell in row.cells:
#                                 if '□' in cell.text and len(cell.text.strip()) < 5:
#                                     cell.text = '☑'
#                                     break
                    
#                     # Mark the appropriate subtopic
#                     for priority_key, priority_data in priority_mapping.items():
#                         if project_type == priority_key:
#                             for subtopic_key, subtopic_text in priority_data['subtopics'].items():
#                                 if subtopic_key.lower() in project_subtopic.lower() and subtopic_text in row_text:
#                                     for cell in row.cells:
#                                         if '□' in cell.text and len(cell.text.strip()) < 5:
#                                             cell.text = '☑'
#                                             break
                
#                 # Handle R&D expenditure table (Table 2)
#                 for row_idx, row in enumerate(rows_list):
#                     row_text = ' '.join([cell.text for cell in row.cells]).strip()
                    
#                     if "For 2022" in row_text and "For 2023" in row_text:
#                         if row_idx + 1 < len(rows_list):
#                             data_row = rows_list[row_idx + 1]
#                             cells = list(data_row.cells)
#                             if len(cells) >= 2:
#                                 if cells[0].text.strip() == "":
#                                     cells[0].text = str(rd_expenditure_2022)
#                                 if cells[1].text.strip() == "":
#                                     cells[1].text = str(rd_expenditure_2023)
                
#                 # Handle research jobs table (Table 3)
#                 for row_idx, row in enumerate(rows_list):
#                     row_text = ' '.join([cell.text for cell in row.cells]).strip()
                    
#                     if "Total number of research jobs created" in row_text and "Research jobs to be created in the project" in row_text:
#                         if row_idx + 1 < len(rows_list):
#                             data_row = rows_list[row_idx + 1]
#                             cells = list(data_row.cells)
#                             if len(cells) >= 3:
#                                 if cells[0].text.strip() == "":
#                                     cells[0].text = str(total_research_jobs)
#                                 if cells[1].text.strip() == "":
#                                     cells[1].text = str(jobs_during_project)
#                                 if cells[2].text.strip() == "":
#                                     cells[2].text = str(jobs_after_project)
            
#             # ========================================================================
#             # 6D. FILL SIGNATURE AREA
#             # ========================================================================
#             # WRITE YOUR SIGNATURE FILLING LOGIC BELOW:
#             for paragraph in doc.paragraphs:
#                 text = paragraph.text
#                 if "title of manager" in text.lower() and "signature" in text.lower() and "name and surname" in text.lower():
#                     new_text = text.replace("(title of manager or person authorised by him/her)", str(manager_title))
#                     new_text = new_text.replace("(name and surname)", str(manager_name))
#                     new_text = new_text.replace("(signature)", "[Signature]")
#                     paragraph.text = new_text
            
#             # ========================================================================
#             # 6E. SAVE AND RETURN
#             # ========================================================================
#             # WRITE YOUR SAVE LOGIC BELOW:
#             doc_bytes = BytesIO()
#             doc.save(doc_bytes)
#             doc_bytes.seek(0)
            
#             return {
#                 "success": True,
#                 "filled_document": doc_bytes,
#                 "filename": f"filled_rd_assessment_{project_type.replace(' ', '_')}.docx",
#                 "preview": f"R&D Assessment: {project_subtopic}, Jobs: {total_research_jobs}, Budget: €{rd_expenditure_2023}"
#             }
            
#         except Exception as e:
#             return {
#                 "success": False,
#                 "error": f"Error filling R&D assessment: {str(e)}"
#             }
#         finally:
#             try:
#                 os.unlink(tmp_path)
#             except:
#                 pass
    
#     # ============================================================================
#     # 7. PASSTHROUGH DOCUMENT PROCESSING (DOCUMENT 4)
#     # ============================================================================
#     def process_passthrough_document(self, uploaded_file):
#         """Process Document 4: Pass-through (no filling, just upload/download)"""
        
#         # WRITE YOUR PASSTHROUGH PROCESSING LOGIC BELOW:
#         try:
#             # Just return the uploaded file as-is
#             file_bytes = BytesIO()
#             file_bytes.write(uploaded_file.getvalue())
#             file_bytes.seek(0)
            
#             # Get original filename or create a default one
#             original_filename = uploaded_file.name if hasattr(uploaded_file, 'name') else "document_4.docx"
#             download_filename = f"processed_{original_filename}"
            
#             return {
#                 "success": True,
#                 "filled_document": file_bytes,
#                 "filename": download_filename,
#                 "preview": f"Document uploaded and ready for download: {original_filename}"
#             }
            
#         except Exception as e:
#             return {
#                 "success": False,
#                 "error": f"Error processing document: {str(e)}"
#             }
    
#     # ============================================================================
#     # 8. PROCESS ALL DOCUMENTS (MAIN ORCHESTRATOR)
#     # ============================================================================
#     def process_all_documents(self, declaration_file, mtep_file, rd_assessment_file, passthrough_file, comprehensive_data):
#         """Process all four documents (3 with filling + 1 pass-through)"""
        
#         # WRITE YOUR ORCHESTRATOR LOGIC BELOW:
#         results = {
#             "declaration_result": None,
#             "mtep_result": None,
#             "rd_assessment_result": None,
#             "passthrough_result": None
#         }
        
#         # Fill declaration document
#         if declaration_file:
#             results["declaration_result"] = self.fill_declaration_form(declaration_file, comprehensive_data)
        
#         # Fill MTEP business plan
#         if mtep_file:
#             results["mtep_result"] = self.fill_mtep_business_plan(mtep_file, comprehensive_data)
        
#         # Fill R&D assessment
#         if rd_assessment_file:
#             results["rd_assessment_result"] = self.fill_rd_assessment_form(rd_assessment_file, comprehensive_data)
        
#         # Process pass-through document (no filling)
#         if passthrough_file:
#             results["passthrough_result"] = self.process_passthrough_document(passthrough_file)
        
#         return results
